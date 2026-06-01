from io import BytesIO
import os
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader

from agents.agent_style.utils import extract_text_from_epub
from .normalizers import decode_text_bytes, estimate_text_tokens, normalize_text, remove_repeated_page_edges
from .types import DocumentSection, ImportWarning, ParsedDocument


SUPPORTED_IMPORT_FORMATS = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".epub": "application/epub+zip",
    ".pdf": "application/pdf",
}

STYLE_IMPORT_FORMATS = tuple(SUPPORTED_IMPORT_FORMATS.keys())
TEXT_SECTION_MIN_CHARS = 20
HEADING_STYLE_RE = re.compile(r"heading\s*[1-9]", re.IGNORECASE)
NOVEL_HEADING_RE = re.compile(
    r"^(?:第\s*[0-9零〇一二三四五六七八九十百千万两]+\s*[章节卷部篇回集]|chapter\s+\d+|prologue|epilogue|序章|终章|番外|楔子|后记)",
    re.IGNORECASE,
)


class UnsupportedImportFormatError(ValueError):
    pass


class ImportTextEmptyError(ValueError):
    pass


def get_supported_formats(usage: str = "general") -> list[str]:
    if usage == "style_analysis":
        return list(STYLE_IMPORT_FORMATS)
    return list(SUPPORTED_IMPORT_FORMATS.keys())


def get_capabilities_payload() -> dict:
    style_formats = get_supported_formats("style_analysis")
    return {
        "success": True,
        "formats": {
            "general": get_supported_formats("general"),
            "style_analysis": style_formats,
        },
        "accept": {
            "general": ",".join(get_supported_formats("general")),
            "style_analysis": ",".join(style_formats),
        },
        "notes": {
            ".pdf": "仅支持带文本层的 PDF，暂不支持扫描件 OCR",
            ".txt": "自动识别常见中文编码",
        },
    }


def parse_uploaded_file(
    file_path: str,
    filename: str | None = None,
    estimate_model: str | None = None,
) -> ParsedDocument:
    suffix = Path(filename or file_path).suffix.lower()
    if suffix not in SUPPORTED_IMPORT_FORMATS:
        raise UnsupportedImportFormatError(f"仅支持 {', '.join(get_supported_formats('general'))} 文件")

    if suffix in {".txt", ".md"}:
        parsed = _parse_text_like_file(file_path, suffix, filename, estimate_model=estimate_model)
    elif suffix == ".docx":
        parsed = _parse_docx_file(file_path, filename, estimate_model=estimate_model)
    elif suffix == ".epub":
        parsed = _parse_epub_file(file_path, filename, estimate_model=estimate_model)
    else:
        parsed = _parse_pdf_file(file_path, filename, estimate_model=estimate_model)

    if not parsed.full_text.strip():
        raise ImportTextEmptyError("无法从文件中提取文本")
    return parsed


def _parse_text_like_file(
    file_path: str,
    suffix: str,
    filename: str | None = None,
    estimate_model: str | None = None,
) -> ParsedDocument:
    with open(file_path, "rb") as f:
        raw_bytes = f.read()
    full_text, encoding_meta = decode_text_bytes(raw_bytes)
    sections = _build_text_sections(full_text, suffix, estimate_model=estimate_model)
    return ParsedDocument(
        filename=filename or os.path.basename(file_path),
        source_format=suffix,
        full_text=full_text,
        sections=sections,
        metadata=encoding_meta,
    )


def _parse_docx_file(
    file_path: str,
    filename: str | None = None,
    estimate_model: str | None = None,
) -> ParsedDocument:
    document = Document(file_path)
    sections: list[DocumentSection] = []
    current_title = ""
    current_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        style_name = getattr(paragraph.style, "name", "") or ""
        is_heading = bool(style_name and HEADING_STYLE_RE.search(style_name))
        if is_heading and current_lines:
            section_text = normalize_text("\n\n".join(current_lines))
            if section_text:
                sections.append(
                    DocumentSection(
                        text=section_text,
                        section_type="heading",
                        title=current_title,
                        metadata={"style_name": style_name},
                        estimated_tokens=estimate_text_tokens(section_text, model=estimate_model),
                    )
                )
            current_lines = []
        if is_heading:
            current_title = text
            if text:
                current_lines.append(text)
            continue
        if text:
            current_lines.append(text)

    if current_lines:
        section_text = normalize_text("\n\n".join(current_lines))
        if section_text:
            sections.append(
                DocumentSection(
                    text=section_text,
                    section_type="heading" if current_title else "paragraph_group",
                    title=current_title,
                    estimated_tokens=estimate_text_tokens(section_text, model=estimate_model),
                )
            )

    full_text = normalize_text("\n\n".join(section.text for section in sections))
    return ParsedDocument(
        filename=filename or os.path.basename(file_path),
        source_format=".docx",
        full_text=full_text,
        sections=sections,
    )


def _parse_epub_file(
    file_path: str,
    filename: str | None = None,
    estimate_model: str | None = None,
) -> ParsedDocument:
    chapter_texts = extract_text_from_epub(file_path, merge_short_chapters=False, min_chunk_size=3000)
    sections = [
        DocumentSection(
            text=normalize_text(chapter_text),
            section_type="chapter",
            title="",
            metadata={"chapter_index": index},
            estimated_tokens=estimate_text_tokens(chapter_text, model=estimate_model),
        )
        for index, chapter_text in enumerate(chapter_texts)
        if normalize_text(chapter_text)
    ]
    full_text = normalize_text("\n\n".join(section.text for section in sections))
    return ParsedDocument(
        filename=filename or os.path.basename(file_path),
        source_format=".epub",
        full_text=full_text,
        sections=sections,
    )


def _parse_pdf_file(
    file_path: str,
    filename: str | None = None,
    estimate_model: str | None = None,
) -> ParsedDocument:
    reader = PdfReader(file_path)
    page_texts: list[str] = []
    for page in reader.pages:
        page_texts.append(normalize_text(page.extract_text() or ""))

    cleaned_pages, removed_edges = remove_repeated_page_edges(page_texts)
    sections = [
        DocumentSection(
            text=page_text,
            section_type="page",
            title=f"第 {index + 1} 页",
            metadata={"page_index": index},
            estimated_tokens=estimate_text_tokens(page_text, model=estimate_model),
        )
        for index, page_text in enumerate(cleaned_pages)
        if page_text
    ]

    warnings: list[ImportWarning] = []
    if not sections:
        warnings.append(ImportWarning(code="pdf_text_layer_missing", message="PDF 未提取到可用文本，可能是扫描件或图片版 PDF"))
    if removed_edges:
        warnings.append(ImportWarning(code="pdf_repeated_edges_removed", message="已自动清理重复页眉页脚"))

    full_text = normalize_text("\n\n".join(section.text for section in sections))
    return ParsedDocument(
        filename=filename or os.path.basename(file_path),
        source_format=".pdf",
        full_text=full_text,
        sections=sections,
        warnings=warnings,
        metadata={"page_count": len(reader.pages)},
    )


def _build_text_sections(
    full_text: str,
    suffix: str,
    estimate_model: str | None = None,
) -> list[DocumentSection]:
    blocks = [normalize_text(block) for block in re.split(r"\n\s*\n+", full_text) if normalize_text(block)]
    sections: list[DocumentSection] = []
    current_lines: list[str] = []
    current_title = ""

    for block in blocks:
        heading = _extract_heading(block, suffix)
        if heading and current_lines:
            section_text = normalize_text("\n\n".join(current_lines))
            if section_text:
                sections.append(
                    DocumentSection(
                        text=section_text,
                        section_type="heading" if current_title else "paragraph_group",
                        title=current_title,
                        estimated_tokens=estimate_text_tokens(section_text, model=estimate_model),
                    )
                )
            current_lines = []
        if heading:
            current_title = heading
        current_lines.append(block)

    if current_lines:
        section_text = normalize_text("\n\n".join(current_lines))
        if section_text:
            sections.append(
                DocumentSection(
                    text=section_text,
                    section_type="heading" if current_title else "paragraph_group",
                    title=current_title,
                    estimated_tokens=estimate_text_tokens(section_text, model=estimate_model),
                )
            )

    return [section for section in sections if len(section.text.strip()) >= TEXT_SECTION_MIN_CHARS]


def _extract_heading(block: str, suffix: str) -> str:
    lines = [line.strip() for line in block.split("\n") if line.strip()]
    if not lines:
        return ""
    first_line = lines[0]
    if suffix == ".md" and first_line.startswith("#"):
        return first_line.lstrip("#").strip()
    if NOVEL_HEADING_RE.match(first_line):
        return first_line
    if len(lines) == 1 and len(first_line) <= 40:
        return first_line
    return ""


def parse_uploaded_bytes(
    raw_bytes: bytes,
    filename: str,
    estimate_model: str | None = None,
) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        tmp_path = BytesIO(raw_bytes)
        full_text, encoding_meta = decode_text_bytes(tmp_path.getvalue())
        sections = _build_text_sections(full_text, suffix, estimate_model=estimate_model)
        return ParsedDocument(
            filename=filename,
            source_format=suffix,
            full_text=full_text,
            sections=sections,
            metadata=encoding_meta,
        )
    raise UnsupportedImportFormatError(f"当前仅支持直接解析 {', '.join(['.txt', '.md'])} 字节内容")
